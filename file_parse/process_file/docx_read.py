# !/usr/bin/env python
# -*- coding:utf-8 -*-　
# @Time : 2023/4/27 09:39 
# @Author : zgf

from typing import Dict
import docx
from utils.common_util import cost_time


@cost_time
async def read_docx_file(file_path) -> Dict[str, object]:
    """
    docx文件读取

    :param doc: 布尔类型，表示是否读取doc文档，默认为False读取docx文档
    :return: 包含标题、作者、最后修改人、最后修改时间、段落和表格等信息的字典对象
    """
    document = docx.Document(file_path)

    paragraphs = [para.text for para in document.paragraphs]

    tables = [
        [
            [cell.text for cell in row.cells]
            for row in table.rows
        ]
        for table in document.tables
    ]

    # 获取元数据
    core_properties = document.core_properties
    title = core_properties.title
    author = core_properties.author
    last_modified_by = core_properties.last_modified_by
    modified = core_properties.modified

    response_data = {
        "file": file_path,
        "metadata": {
            "title": title,
            "author": author,
            "last_modified_by": last_modified_by,
            "modified": modified
        },
        "content": {
            "paragraph": paragraphs,
            "tables": tables,
        },
    }
    return response_data
